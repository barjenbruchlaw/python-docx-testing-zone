from pydoc import plain
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

county_name = "Jackson"
plaintiff_name = "FirstKey Homes, as agent for CSMA BLT, LLC"
defendant1_name = "John Smith"
defendant2_name = "Mary Smith"
defendant_address1 = "12345 Main Street"
defendant_address2 = "101"
defendant_city = "Kansas City"
defendant_state = "MO"
defendant_zip = "64114"
case_number = "22XX-CV000000"
division = "99"

caption_first_line = f'In the circuit court for {county_name} County, Missouri'

caption_first_line_formatted = caption_first_line.upper()

plaintiff_name_formatted = plaintiff_name.upper()
defendant1_name_formatted = defendant1_name.upper()
defendant2_name_formatted = defendant2_name.upper()

case_number_text = "Case no. "+case_number
division_text = "Division "+division

jurisdiction_line = document.add_paragraph()

jurisdiction_line.style = document.styles['Normal']

jurisdiction_line_format = jurisdiction_line.paragraph_format
jurisdiction_line_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

jurisdiction_line.add_run(caption_first_line_formatted).bold=True

document.add_paragraph()

caption_table = document.add_table(rows=3, cols=2)

first_row = caption_table.rows[0].cells
second_row = caption_table.rows[1].cells
third_row = caption_table.rows[2].cells

plaintiff_area = first_row[0]

plaintiff_area.text = plaintiff_name_formatted + ",\n \nPlaintiff,"

versus = second_row[0]

versus.text = "\n  v."

case_number_area = second_row[1]

case_number_area.text = case_number_text + "\n\n" + division_text

defendant_area = third_row[0]

defendant_area.text = defendant1_name_formatted+"\n"+defendant2_name_formatted+"\nJOHN DOE and/or MARY ROE\n"+defendant_address1+" #"+defendant_address2+"\n"+defendant_city+", "+defendant_state+" "+defendant_zip+"\n\nDefendants."

document.save('caption.docx')

